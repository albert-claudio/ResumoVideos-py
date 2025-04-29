import os
from pytube import YouTube
import whisper
from transformers import pipeline
from docx import Document
from fpdf import FPDF

# === CONFIGURAÇÕES ===
# Modelo Whisper para transcrição (baixar automaticamente na primeira vez)
WHISPER_MODEL = "base"
# Modelo de resumo local (Hugging Face)
SUMMARIZATION_MODEL = "sshleifer/distilbart-cnn-12-6"
# Tamanho máximo de caracteres por chunk
MAX_CHARS_PER_CHUNK = 1200

# === CARREGANDO MODELOS ===
print("🔊 Carregando Whisper para transcrição profunda...")
whisper_model = whisper.load_model(WHISPER_MODEL)
print("📝 Carregando pipeline de resumo...")
summarizer = pipeline("summarization", model=SUMMARIZATION_MODEL)

# === FUNÇÕES AUXILIARES ===

def baixar_audio_youtube(url: str, output_path: str = 'audio.mp4') -> str:
    """
    Faz download do áudio do YouTube via pytube.
    """
    yt = YouTube(url)
    stream = yt.streams.filter(only_audio=True).first()
    out_file = stream.download(filename=output_path)
    return out_file


def obter_transcricao(url: str) -> str:
    """
    Obtém transcrição completa usando Whisper.
    """
    audio_file = baixar_audio_youtube(url)
    print("⏳ Transcrevendo áudio com Whisper (pode levar alguns minutos)...")
    result = whisper_model.transcribe(audio_file, language='pt')
    return result.get('text', '').strip()


def dividir_em_chunks(text: str, max_chars: int) -> list[str]:
    """
    Divide texto em pedaços de até max_chars, quebrando em espaços.
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chars
        if end >= len(text):
            chunks.append(text[start:].strip())
            break
        split = text.rfind(' ', start, end)
        if split <= start:
            split = end
        chunks.append(text[start:split].strip())
        start = split
    return chunks


def resumir_texto(text: str) -> str:
    """
    Resume o texto em vários chunks e concatena os mini-resumos.
    """
    print("📚 Dividindo texto em chunks...")
    chunks = dividir_em_chunks(text, MAX_CHARS_PER_CHUNK)
    sumarios = []
    for i, chunk in enumerate(chunks, 1):
        print(f"🔹 Resumindo chunk {i}/{len(chunks)}...")
        out = summarizer(chunk, max_length=150, min_length=40, do_sample=False)
        sumarios.append(out[0]['summary_text'])
    print("🔄 Combinando resumos...")
    return ' '.join(sumarios)


def exportar_word(text: str, filename: str):
    doc = Document()
    doc.add_heading('Resumo de Vídeo', 0)
    doc.add_paragraph(text)
    doc.save(filename)


def exportar_pdf(text: str, filename: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font('Arial', size=12)
    for line in text.split('\n'):
        pdf.multi_cell(0, 8, line)
    pdf.output(filename)


def main():
    url = input("📺 Cole a URL do vídeo do YouTube: ").strip()
    if not url:
        print("⚠️  URL vazia. Encerrando.")
        return

    print('1. Obtendo transcrição...')
    texto = obter_transcricao(url)
    if not texto:
        print('⚠️  Não foi possível transcrever o vídeo.')
        return

    print('2. Resumindo conteúdo...')
    resumo = resumir_texto(texto)

    word_file = input("Nome do arquivo .docx de saída (Enter para 'resumo.docx'): ").strip() or 'resumo.docx'
    pdf_file = input("Nome do arquivo .pdf de saída (Enter para 'resumo.pdf'): ").strip() or 'resumo.pdf'

    print(f"3. Exportando para Word: {word_file}")
    exportar_word(resumo, word_file)
    print(f"4. Exportando para PDF: {pdf_file}")
    exportar_pdf(resumo, pdf_file)

    print('✅ Processo concluído!')

if __name__ == '__main__':
    main()
